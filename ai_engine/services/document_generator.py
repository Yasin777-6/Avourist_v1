"""
Document Generator Service
Generates .docx files for petitions and contracts
"""

import os
import logging
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Generate professional legal documents in .docx format"""
    
    def __init__(self):
        self.temp_dir = Path('temp_documents')
        self.temp_dir.mkdir(exist_ok=True)
    
    def generate_petition_docx(self, petition_text: str, client_name: str = None) -> str:
        """
        Generate a .docx file from petition text
        
        Args:
            petition_text: The petition content
            client_name: Client name for filename
            
        Returns:
            Path to generated .docx file
        """
        try:
            # Strip HTML tags from petition text
            import re
            petition_text = re.sub(r'<[^>]+>', '', petition_text)
            
            # Remove separator lines (────────)
            petition_text = re.sub(r'─+', '', petition_text)
            
            # Create document
            doc = Document()
            
            # Set margins (2.5cm all sides - standard for Russian legal docs)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.2)
                section.right_margin = Inches(0.8)
            
            # Parse petition text
            lines = petition_text.strip().split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Detect line type and format accordingly
                if 'ХОДАТАЙСТВО' in line.upper() and len(line) < 50:
                    # Title
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(14)
                
                elif line.startswith('В ') and 'суд' in line.lower():
                    # Court address (right-aligned)
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.runs[0].font.size = Pt(12)
                
                elif line.startswith('От:') or line.startswith('от '):
                    # From (right-aligned)
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.runs[0].font.size = Pt(12)
                
                elif line.lower().startswith('адрес:'):
                    # Address (right-aligned)
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.runs[0].font.size = Pt(12)
                
                elif line.startswith('ПРОШУ:'):
                    # Request section
                    p = doc.add_paragraph(line)
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(12)
                
                elif line.startswith('Дата:') or line.startswith('Подпись:'):
                    # Signature block
                    p = doc.add_paragraph(line)
                    p.runs[0].font.size = Pt(12)
                
                elif line.startswith('Приложение:'):
                    # Attachments
                    p = doc.add_paragraph(line)
                    p.runs[0].font.size = Pt(11)
                
                else:
                    # Regular paragraph
                    p = doc.add_paragraph(line)
                    p.runs[0].font.size = Pt(12)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Generate filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            if client_name:
                safe_name = "".join(c for c in client_name if c.isalnum() or c in (' ', '_')).strip()
                filename = f"Ходатайство_{safe_name}_{timestamp}.docx"
            else:
                filename = f"Ходатайство_{timestamp}.docx"
            
            filepath = self.temp_dir / filename
            
            # Save document
            doc.save(str(filepath))
            
            logger.info(f"Generated petition document: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generating petition document: {e}")
            raise
    
    def generate_contract_docx(self, contract_text: str, client_name: str = None) -> str:
        """
        Generate a .docx file from contract text
        
        Args:
            contract_text: The contract content
            client_name: Client name for filename
            
        Returns:
            Path to generated .docx file
        """
        try:
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.2)
                section.right_margin = Inches(0.8)
            
            # Parse contract text
            lines = contract_text.strip().split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                if 'ДОГОВОР' in line.upper() and len(line) < 100:
                    # Title
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(14)
                
                elif line.endswith(':') and len(line) < 50:
                    # Section headers
                    p = doc.add_paragraph(line)
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(12)
                
                else:
                    # Regular paragraph
                    p = doc.add_paragraph(line)
                    p.runs[0].font.size = Pt(12)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Generate filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            if client_name:
                safe_name = "".join(c for c in client_name if c.isalnum() or c in (' ', '_')).strip()
                filename = f"Договор_{safe_name}_{timestamp}.docx"
            else:
                filename = f"Договор_{timestamp}.docx"
            
            filepath = self.temp_dir / filename
            
            # Save document
            doc.save(str(filepath))
            
            logger.info(f"Generated contract document: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generating contract document: {e}")
            raise
    
    def generate_and_send_petition(self, lead, params: str) -> str:
        """
        Generate petition document and send it to user via Telegram
        
        Args:
            lead: Lead object
            params: Pipe-separated parameters (type|name|address|court|reason|additional)
            
        Returns:
            Response message to user
        """
        try:
            logger.info(f"Generating petition for lead {lead.telegram_id} with params: {params}")
            
            # Parse parameters
            parts = params.split('|')
            if len(parts) < 4:
                logger.error(f"Insufficient parameters: {parts}")
                return "❌ Ошибка: недостаточно данных для создания ходатайства."
            
            petition_type = parts[0].strip()
            client_name = parts[1].strip() if len(parts) > 1 else lead.first_name or "Клиент"
            address = parts[2].strip() if len(parts) > 2 else ""
            court = parts[3].strip() if len(parts) > 3 else ""
            reason = parts[4].strip() if len(parts) > 4 else ""
            additional = parts[5].strip() if len(parts) > 5 else ""
            
            # Generate petition text based on type
            petition_text = self._generate_petition_text(
                petition_type, client_name, address, court, reason, additional
            )
            
            # Generate .docx file
            file_path = self.generate_petition_docx(petition_text, client_name)
            
            # Send to Telegram
            self._send_petition_to_telegram(lead.telegram_id, file_path, petition_type)
            
            # Return success message
            return self._get_petition_success_message(petition_type)
            
        except Exception as e:
            logger.error(f"Error generating petition: {e}")
            logger.exception("Petition generation error:")
            return "❌ Произошла ошибка при создании ходатайства. Попробуйте еще раз или обратитесь к менеджеру."
    
    def _generate_petition_text(self, petition_type: str, client_name: str, 
                                address: str, court: str, reason: str, additional: str) -> str:
        """Generate petition text based on type and parameters"""
        
        today = datetime.now().strftime('%d.%m.%Y')
        
        # Common header
        header = f"""В {court}
от {client_name}
Адрес: {address}

"""
        
        # Generate petition based on type
        if petition_type == 'postpone_hearing':
            title = "ХОДАТАЙСТВО\nо переносе судебного заседания"
            body = f"""
По делу об административном правонарушении в отношении {client_name}.

{additional}

В связи с {reason}, я не могу явиться в назначенное судебное заседание.

На основании изложенного, руководствуясь ст. 24.4 КоАП РФ,

ПРОШУ:
1. Перенести рассмотрение дела об административном правонарушении на более позднюю дату.

Дата: {today}
Подпись: ___________ ({client_name})
"""
        
        elif petition_type == 'evidence_attachment':
            title = "ХОДАТАЙСТВО\nо приобщении доказательств"
            body = f"""
По делу об административном правонарушении в отношении {client_name}.

{additional}

Прошу приобщить к материалам дела следующие доказательства:
{reason}

Указанные доказательства имеют существенное значение для правильного рассмотрения дела и установления всех обстоятельств.

На основании изложенного, руководствуясь ст. 24.4, 26.2 КоАП РФ,

ПРОШУ:
1. Приобщить к материалам дела представленные доказательства.
2. Приобщить к материалам дела видеозапись на электронном носителе.

Приложение:
1. Электронный носитель (флешка) с видеозаписью - 1 шт.

Дата: {today}
Подпись: ___________ ({client_name})
"""
        
        elif petition_type == 'materials_access':
            title = "ХОДАТАЙСТВО\nоб ознакомлении с материалами дела"
            body = f"""
По делу об административном правонарушении в отношении {client_name}.

В соответствии со ст. 25.1 КоАП РФ лицо, в отношении которого ведется производство по делу об административном правонарушении, вправе знакомиться со всеми материалами дела.

{reason}

На основании изложенного, руководствуясь ст. 24.4, 25.1 КоАП РФ,

ПРОШУ:
1. Предоставить возможность ознакомиться с материалами дела об административном правонарушении.
2. Предоставить копии материалов дела (протоколы, схемы, фотографии и т.д.).

Дата: {today}
Подпись: ___________ ({client_name})
"""
        
        elif petition_type == 'expertise':
            title = "ХОДАТАЙСТВО\nо назначении экспертизы"
            body = f"""
По делу об административном правонарушении в отношении {client_name}.

{additional}

Считаю необходимым назначение независимой экспертизы по следующим основаниям:
{reason}

Для установления объективной истины и правильного разрешения дела необходимо проведение экспертизы.

На основании изложенного, руководствуясь ст. 24.4, 26.4 КоАП РФ,

ПРОШУ:
1. Назначить независимую экспертизу по делу.
2. Поставить на разрешение эксперта вопросы об исправности прибора и достоверности его показаний.

Дата: {today}
Подпись: ___________ ({client_name})
"""
        
        elif petition_type == 'witness_summon':
            title = "ХОДАТАЙСТВО\nо вызове свидетелей"
            body = f"""
По делу об административном правонарушении в отношении {client_name}.

{additional}

Для установления всех обстоятельств дела прошу вызвать в судебное заседание свидетелей:
{reason}

Показания указанных свидетелей имеют существенное значение для правильного рассмотрения дела.

На основании изложенного, руководствуясь ст. 24.4, 25.6 КоАП РФ,

ПРОШУ:
1. Вызвать в судебное заседание указанных свидетелей.
2. Допросить свидетелей по обстоятельствам дела.

Дата: {today}
Подпись: ___________ ({client_name})
"""
        
        elif petition_type == 'license_return':
            title = "ХОДАТАЙСТВО\nо возврате водительского удостоверения"
            body = f"""
{additional}

Срок лишения права управления транспортными средствами истек {reason}.

Медицинскую справку получил, теоретический экзамен в ГИБДД сдал успешно.

На основании изложенного, руководствуясь ст. 32.6, 32.7 КоАП РФ,

ПРОШУ:
1. Вернуть водительское удостоверение.

Приложение:
1. Медицинская справка
2. Квитанция об оплате штрафов

Дата: {today}
Подпись: ___________ ({client_name})
"""
        
        else:
            # Generic petition
            title = "ХОДАТАЙСТВО"
            body = f"""
По делу об административном правонарушении в отношении {client_name}.

{reason}

{additional}

На основании изложенного, руководствуясь ст. 24.4 КоАП РФ,

ПРОШУ:
1. Удовлетворить настоящее ходатайство.

Дата: {today}
Подпись: ___________ ({client_name})
"""
        
        return header + title + body
    
    def _send_petition_to_telegram(self, telegram_id: int, file_path: str, petition_type: str):
        """Send petition document to Telegram user"""
        import requests
        from django.conf import settings
        
        try:
            logger.info(f"Sending petition to Telegram user {telegram_id}")
            
            url = f"https://api.telegram.org/bot{settings.TELEGRAM_BOT_TOKEN}/sendDocument"
            
            with open(file_path, "rb") as file:
                files = {"document": file}
                data = {
                    "chat_id": telegram_id,
                    "caption": "📄 Ваше ходатайство готово",
                    "parse_mode": "HTML",
                }
                
                response = requests.post(url, data=data, files=files, timeout=30)
                response.raise_for_status()
                logger.info(f"✅ Petition sent successfully to {telegram_id}")
                
        except Exception as e:
            logger.error(f"❌ Failed to send petition to {telegram_id}: {str(e)}")
            logger.exception("Telegram send error:")
            raise
    
    def _get_petition_success_message(self, petition_type: str) -> str:
        """Get success message based on petition type"""
        
        base_message = "✅ Ходатайство готово! Отправил вам файл .docx\n\n📋 <b>Как подать:</b>\n"
        
        if petition_type == 'postpone_hearing':
            return base_message + """1. Скачайте файл выше
2. Распечатайте в 2 экземплярах
3. Подпишите оба экземпляра
4. Подайте в канцелярию суда СЕГОДНЯ
5. Один экземпляр с отметкой оставьте себе

⚠️ <b>Важно:</b> Также заявите ходатайство устно в начале заседания!"""
        
        elif petition_type == 'evidence_attachment':
            return base_message + """1. Скачайте файл выше
2. Распечатайте в 2 экземплярах
3. Запишите видео на флешку
4. Подайте в канцелярию суда до заседания
5. На второй экземпляр попросите поставить отметку о принятии

⚠️ <b>Важно:</b> Флешку приложите к ходатайству!"""
        
        elif petition_type == 'materials_access':
            return base_message + """1. Скачайте файл выше
2. Распечатайте в 2 экземплярах
3. Подпишите
4. Подайте в канцелярию суда
5. На втором экземпляре попросите поставить отметку

📌 Обычно материалы предоставляют в течение 1-3 дней."""
        
        elif petition_type == 'expertise':
            return base_message + """1. Скачайте файл выше
2. Распечатайте в 2 экземплярах
3. Подпишите
4. Подайте в канцелярию суда до заседания
5. Заявите ходатайство в суде устно

⚠️ <b>Важно:</b> Суд может отказать, если нет оснований для экспертизы."""
        
        elif petition_type == 'witness_summon':
            return base_message + """1. Скачайте файл выше
2. Распечатайте в 2 экземплярах
3. Подпишите
4. Подайте в канцелярию суда заранее
5. Укажите контакты свидетелей

📌 Суд вызовет свидетелей повесткой."""
        
        elif petition_type == 'license_return':
            return base_message + """1. Скачайте файл выше
2. Распечатайте
3. Подпишите
4. Приложите медсправку и квитанции об оплате штрафов
5. Подайте в ГИБДД, где хранятся права

📌 Права вернут в течение 1-3 дней."""
        
        else:
            return base_message + """1. Скачайте файл выше
2. Распечатайте в 2 экземплярах
3. Подпишите
4. Подайте в канцелярию суда
5. Один экземпляр с отметкой оставьте себе"""
    
    def cleanup_old_documents(self, days: int = 1):
        """Delete documents older than specified days"""
        try:
            import time
            now = time.time()
            cutoff = now - (days * 86400)
            
            for file in self.temp_dir.glob('*.docx'):
                if file.stat().st_mtime < cutoff:
                    file.unlink()
                    logger.info(f"Deleted old document: {file}")
        except Exception as e:
            logger.error(f"Error cleaning up documents: {e}")
