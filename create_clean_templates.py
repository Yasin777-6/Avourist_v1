"""
Create clean DOCX templates for each instance type.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_template(instance: str, output_path: str):
    """Create a clean DOCX template."""
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Договор № _____ об оказании юридических услуг")
    run.bold = True
    run.font.size = Pt(14)
    
    # Date and location
    date_para = doc.add_paragraph()
    date_para.add_run("г. Москва                                                    ")
    date_run = date_para.add_run("28 апреля 2024 г")
    date_run.font.color.rgb = RGBColor(255, 0, 0)
    
    # Parties
    parties = doc.add_paragraph()
    name_run = parties.add_run("Тытюк Александр Михайлович")
    name_run.font.color.rgb = RGBColor(255, 0, 0)
    parties.add_run(" (именуемый в дальнейшем «Заказчик») и Автономная некоммерческая организация по предоставлению услуг в области права «Первое Арбитражное Учреждение» в лице директора Шельмина Евгения Васильевича, (именуемый в дальнейшем «Исполнитель»), с другой стороны, именуемые в дальнейшем «Стороны», заключили настоящий Договор о нижеследующем:")
    
    # Section 1
    doc.add_heading("1. Предмет договора", level=2)
    section1 = doc.add_paragraph()
    section1.add_run("1.1 Заказчик поручает, а Исполнитель принимает на себя обязательства оказать юридические услуги: провести правовой анализ документов (материалов дела), ")
    case_desc_run = section1.add_run("подготовка ответа на требование каршеринга, подготовка претензии, подготовка Отзыва на исковое заявления, ответ на претензию, заявление в соответствующие органы")
    case_desc_run.font.color.rgb = RGBColor(255, 0, 0)
    section1.add_run(". Исполнитель представляет интересы Заказчика в суде и в иных органах по нотариальной доверенности без личного присутствия.")
    
    # Section 5 - Pricing
    doc.add_heading("5. Стоимость услуг и порядок расчетов", level=2)
    
    pricing_para = doc.add_paragraph()
    pricing_para.add_run("5.1. Стоимость услуг по договору в соответствии с п.1.1. настоящего Договора составляет ")
    total_run = pricing_para.add_run("_________(______________)  рублей")
    total_run.font.color.rgb = RGBColor(255, 0, 0)
    pricing_para.add_run(". Оплата производится в следующем порядке: ")
    prep_run = pricing_para.add_run("____________ (_________ тысяч) рублей")
    prep_run.font.color.rgb = RGBColor(255, 0, 0)
    pricing_para.add_run(" не позднее одного дня с даты подписания настоящего Договора Сторонами. ")
    success_run = pricing_para.add_run("__________________ (___________ тысяч) рублей")
    success_run.font.color.rgb = RGBColor(255, 0, 0)
    pricing_para.add_run(" исключительно при положительном решении в соответствии с п.8.8 настоящего Договора.")
    
    doc.add_paragraph()
    docs_para = doc.add_paragraph()
    docs_para.add_run("5.1.1 Подготовка документов в соответствии с п.1.1. настоящего Договора составляет ")
    docs_run = docs_para.add_run("______ (___________ тысяч) рублей")
    docs_run.font.color.rgb = RGBColor(255, 0, 0)
    docs_para.add_run(".")
    
    # Signatures section
    doc.add_heading("9. Адреса, реквизиты и подписи Сторон:", level=2)
    
    # Create table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Left column
    left_cell = table.rows[0].cells[0]
    left_cell.text = "Исполнитель\n\nАНО «Первое Арбитражное Учреждение»\n\n_______________________\nДиректор Шельмин Е.В."
    
    # Right column
    right_cell = table.rows[0].cells[1]
    client_info = right_cell.add_paragraph()
    
    client_name_run = client_info.add_run("Тытюк Александр Михайлович")
    client_name_run.font.color.rgb = RGBColor(255, 0, 0)
    
    client_info.add_run("\n\nОБЯЗАТЕЛЬНО !!!!!!!!!!!!!!!!!!!\n")
    
    birth_date_run = client_info.add_run("Дата/ месяц/ год рождения")
    birth_date_run.font.color.rgb = RGBColor(255, 0, 0)
    client_info.add_run("\n")
    
    birth_place_run = client_info.add_run("Место рождения")
    birth_place_run.font.color.rgb = RGBColor(255, 0, 0)
    client_info.add_run("\n")
    
    passport_run = client_info.add_run("Паспорт Серия_____ Номер___________")
    passport_run.font.color.rgb = RGBColor(255, 0, 0)
    client_info.add_run("\n\n")
    
    address_run = client_info.add_run("Зарегистрирован: _____________________")
    address_run.font.color.rgb = RGBColor(255, 0, 0)
    client_info.add_run("\n\n")
    
    phone_run = client_info.add_run("Тел. _________________________")
    phone_run.font.color.rgb = RGBColor(255, 0, 0)
    client_info.add_run("\n\n")
    
    email_run = client_info.add_run("Е-mail______________________________")
    email_run.font.color.rgb = RGBColor(255, 0, 0)
    client_info.add_run("\n\n")
    
    signature_run = client_info.add_run("_______________________________________\n(Фамилия И.О прописью / подпись)")
    signature_run.font.color.rgb = RGBColor(255, 0, 0)
    
    # Save
    doc.save(output_path)
    print(f"✓ Created: {output_path}")

# Create templates for all instances
contracts_dir = Path("contracts")
contracts_dir.mkdir(exist_ok=True)

templates = [
    ("1", "Договор_БЕЗ_ДОВЕРЕННОСТИ_на_1_инстанцию.docx"),
    ("2", "Договор_БЕЗ_ДОВЕРЕННОСТИ_на_2_инстанции.docx"),
    ("3", "Договор_БЕЗ_ДОВЕРЕННОСТИ_на_3_инстанции.docx"),
    ("4", "Договор_БЕЗ_ДОВЕРЕННОСТИ_на_4_инстанции.docx"),
]

for instance, filename in templates:
    output_path = contracts_dir / filename
    create_template(instance, str(output_path))

print("\n✅ All templates created successfully!")
