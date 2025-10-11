"""
DOCX-based contract filler - searches for placeholder text and replaces with actual data.
Much simpler and more reliable than PDF manipulation.
"""
import logging
from pathlib import Path
from typing import Dict
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt

logger = logging.getLogger(__name__)


class DOCXFiller:
    """
    Simple DOCX filler that searches for placeholder text and replaces it.
    """
    
    def __init__(self):
        pass
    
    def fill_docx(self, template_path: str, data: Dict, output_path: str) -> bytes:
        """
        Fill DOCX template by replacing placeholder text with actual data.
        
        Args:
            template_path: Path to DOCX template
            data: Dictionary with contract data
            output_path: Path to save filled DOCX
            
        Returns:
            DOCX bytes
        """
        logger.info(f"Opening template: {template_path}")
        
        # Normalize data
        data = self._normalize_data(data)
        
        # For .doc files, try to open them directly (they might be DOCX with wrong extension)
        try:
            doc = Document(template_path)
        except Exception as e:
            logger.error(f"Failed to open template: {e}")
            # If it fails, just copy the template and return it
            import shutil
            shutil.copy(template_path, output_path)
            with open(output_path, 'rb') as f:
                return f.read()
        
        # Define replacements based on your field list
        replacements = self._build_replacements(data)
        
        # Log what we're replacing
        logger.info(f"=== DOCX REPLACEMENTS ===")
        logger.info(f"Input data:")
        logger.info(f"  Passport: {data.get('client_passport_series', 'MISSING')} / {data.get('client_passport_number', 'MISSING')}")
        logger.info(f"  Case article: {data.get('case_article', 'MISSING')}")
        logger.info(f"  Total amount: {data.get('total_amount', 'MISSING')}")
        logger.info(f"  Prepayment: {data.get('prepayment', 'MISSING')}")
        logger.info(f"  Success fee: {data.get('success_fee', 'MISSING')}")
        logger.info(f"  Docs fee: {data.get('docs_prep_fee', 'MISSING')}")
        logger.info(f"  Payment terms desc: {data.get('payment_terms_description', 'MISSING')}")
        logger.info(f"Total replacements: {len(replacements)}")
        logger.info(f"First 15 replacements:")
        for key, value in list(replacements.items())[:15]:
            logger.info(f"  '{key[:50]}...' → '{str(value)[:50]}...'")
        
        # Count replacements made
        replacements_made = 0
        
        # Replace text in paragraphs
        for paragraph in doc.paragraphs:
            if self._replace_in_paragraph(paragraph, replacements):
                replacements_made += 1
        
        # Replace text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if self._replace_in_paragraph(paragraph, replacements):
                            replacements_made += 1
        
        logger.info(f"Total paragraphs/cells modified: {replacements_made}")
        
        # Save filled document
        doc.save(output_path)
        
        # Read and return bytes
        with open(output_path, 'rb') as f:
            docx_bytes = f.read()
        
        logger.info(f"Filled DOCX saved to: {output_path} ({len(docx_bytes)} bytes)")
        return docx_bytes
    
    def _build_replacements(self, data: Dict) -> Dict[str, str]:
        """
        Build dictionary of placeholder text → replacement text.
        Based on the user's sample data.
        """
        replacements = {}
        
        # Contract number: "Договор № 1-Б/24"
        contract_number = data.get('contract_number', '')
        if contract_number:
            replacements["Договор № 1-Б/24"] = f"Договор № {contract_number}"
            replacements["№ 1-Б/24"] = f"№ {contract_number}"
        
        # Date field: "28 апреля 2024 г"
        contract_date = data.get('contract_date', datetime.now().strftime('%d.%m.%Y'))
        replacements["28 апреля 2024 г"] = contract_date
        replacements["г. Москва28 апреля 2024 г"] = f"г. Москва {contract_date}"
        
        # Client name: "Тытюк Александр Михайлович"
        client_name = data.get('client_full_name', '')
        if client_name:
            replacements["Тытюк  Александр  Михайлович"] = client_name
            replacements["Тытюк Александр Михайлович"] = client_name
        
        # Director name (Executor side)
        director_name = data.get('director_name', 'Шельмина Евгения Васильевича')
        replacements["Шельмина Евгения Васильевича"] = director_name
        
        # Pricing fields
        total_amount = data.get('total_amount')
        if total_amount:
            amount_digits = f"{int(total_amount):,}".replace(",", " ")
            amount_words = self._amount_to_words(int(total_amount))
            replacements["_________(______________) рублей"] = f"{amount_digits} ({amount_words}) рублей"
            replacements["_________(______________)  рублей"] = f"{amount_digits} ({amount_words}) рублей"
            replacements["25 000 (двадцать пять тысяч) рублей"] = f"{amount_digits} ({amount_words}) рублей"
        
        prepayment = data.get('prepayment')
        if prepayment:
            prep_digits = f"{int(prepayment):,}".replace(",", " ")
            prep_words = self._amount_to_words(int(prepayment))
            replacements["____________ (_________ тысяч)рублей"] = f"{prep_digits} ({prep_words}) рублей"
            replacements["____________ (_________ тысяч) рублей"] = f"{prep_digits} ({prep_words}) рублей"
            replacements["15 000 (пятнадцать тысяч) рублей"] = f"{prep_digits} ({prep_words}) рублей"
        
        success_fee = data.get('success_fee')
        if success_fee:
            fee_digits = f"{int(success_fee):,}".replace(",", " ")
            fee_words = self._amount_to_words(int(success_fee))
            replacements["__________________ (___________ тысяч) рублей"] = f"{fee_digits} ({fee_words}) рублей"
            replacements["__________________(___________ тысяч) рублей"] = f"{fee_digits} ({fee_words}) рублей"
            replacements["10 000 (десять тысяч) рублей"] = f"{fee_digits} ({fee_words}) рублей"
        
        docs_fee = data.get('docs_prep_fee')
        if docs_fee:
            docs_digits = f"{int(docs_fee):,}".replace(",", " ")
            docs_words = self._amount_to_words(int(docs_fee))
            replacements["______(___________ тысяч) рублей"] = f"{docs_digits} ({docs_words}) рублей"
        
        # Client details section
        birth_date = data.get('birth_date', '')
        if birth_date:
            replacements["Дата/ месяц/ год рождения"] = f"Дата/ месяц/ год рождения: {birth_date}"
        
        birth_place = data.get('birth_place', '')
        if birth_place:
            replacements["Место рождения"] = f"Место рождения: {birth_place}"
        
        passport_series = data.get('client_passport_series', '')
        passport_number = data.get('client_passport_number', '')
        if passport_series and passport_number:
            # Multiple possible variations of passport placeholder
            passport_variations = [
                "Паспорт Серия_____ Номер___________",
                "Паспорт Серия_____  Номер___________",  # Note: 2 spaces
                "Паспорт Серия_____ Номер__________",   # 10 underscores
                "Серия_____ Номер___________",
                "Серия_____  Номер___________",          # Note: 2 spaces
            ]
            passport_text = f"Паспорт Серия {passport_series} Номер {passport_number}"
            series_text = f"Серия {passport_series} Номер {passport_number}"
            
            for variation in passport_variations:
                if "Паспорт" in variation:
                    replacements[variation] = passport_text
                else:
                    replacements[variation] = series_text
        
        address = data.get('client_address', '')
        if address:
            replacements["Зарегистрирован: _____________________"] = f"Зарегистрирован: {address}"
        
        phone = data.get('client_phone', '')
        if phone:
            replacements["Тел. _________________________"] = f"Тел. {phone}"
        
        email = data.get('email', '')
        if email:
            replacements["Е-mail______________________________"] = f"Е-mail: {email}"
        
        # Case article - replace in default descriptions and standalone
        case_article = data.get('case_article', '')
        if case_article:
            # Replace standalone article references
            replacements["ч.1 ст.12.8 КоАП РФ"] = case_article
            replacements["ч.1 ст.12.8 КоАП"] = case_article
            
            # Replace article in default descriptions
            default_descs_with_article = [
                "подготовитьходатайство на получение материалов дела, ходатайство о привлечения к делу защитника, подготовка жалобы на постановление по делу об административном правонарушении по ч.1 ст.12.8 КоАП РФ",
                "получение материалов дела, ходатайство о привлечении к делу защитника, ходатайство о переводе дела по месту жительства, подготовка письменного объяснения лица по делу об административном правонарушении по ч.1 ст.12.8 КоАП РФ",
                "получение материалов дела, ходатайство о привлечении к делу защитника, ходатайство о переводе дела по месту жительства, подготовка письменного объяснения лица по делу об административном правонарушении по ч.1 ст.12.8 КоАП РФ."
            ]
            for default_desc in default_descs_with_article:
                # Replace with actual article
                new_desc = default_desc.replace("ч.1 ст.12.8 КоАП РФ", case_article)
                replacements[default_desc] = new_desc
        
        # Case description - from user's sample
        case_desc = data.get('case_description', '')
        if case_desc:
            # Multiple possible default descriptions
            default_descs = [
                "провести правовой анализ документов (материалов дела),подготовка ответа на требование каршеринга, подготовка претензии, подготовка Отзыва на исковое заявления, ответ на претензию,заявление в соответствующие органы",
                "подготовка ответа на требование каршеринга, подготовка претензии, подготовка Отзыва на исковое заявления, ответ на претензию,заявление в соответствующие органы",
                "получение материалов дела, ходатайство о привлечении к делу защитника, ходатайство о переводе дела по месту жительства, подготовка письменного объяснения лица по делу об административном правонарушении"
            ]
            for default_desc in default_descs:
                replacements[default_desc] = case_desc
        
        # Payment terms description
        payment_terms_desc = data.get('payment_terms_description', '')
        if payment_terms_desc:
            # Replace default payment terms placeholders
            default_payment_terms = [
                "__% предоплата, __% после положительного решения",
                "___% предоплата, ___% после положительного решения",
                "50% предоплата, 50% после положительного решения"
            ]
            for default_term in default_payment_terms:
                replacements[default_term] = payment_terms_desc
        
        # Representation clause
        representation_clause = data.get('representation_clause', '')
        if representation_clause:
            default_clause = "Исполнитель представляет интересы Заказчика в судеи в иных органах по нотариальной доверенности без личного присутствия."
            replacements[default_clause] = representation_clause
        
        return replacements
    
    def _replace_in_paragraph(self, paragraph, replacements: Dict[str, str]) -> bool:
        """Replace placeholder text in a paragraph while preserving formatting.
        Returns True if any replacement was made."""
        full_text = paragraph.text
        original_text = full_text
        
        # Check if any replacement is needed
        for placeholder, replacement in replacements.items():
            if placeholder in full_text:
                # Replace the text
                full_text = full_text.replace(placeholder, replacement)
        
        # If text was modified, update the paragraph
        if full_text != original_text:
            # Clear existing runs
            for run in paragraph.runs:
                run.text = ""
            
            # Add new text
            if paragraph.runs:
                paragraph.runs[0].text = full_text
            else:
                paragraph.add_run(full_text)
            return True
        return False
    
    def _amount_to_words(self, amount: int) -> str:
        """Convert amount to Russian words."""
        try:
            from num2words import num2words
            return num2words(int(amount), lang="ru")
        except Exception:
            return str(amount)
    
    def _normalize_data(self, data: Dict) -> Dict:
        """Normalize incoming data keys."""
        def first(*keys):
            for k in keys:
                v = data.get(k)
                if v not in (None, ""):
                    return v
            return None

        norm = dict(data)
        norm.setdefault("client_full_name", first("client_full_name", "full_name", "fio", "client_name") or "")
        norm.setdefault("client_passport_series", first("client_passport_series", "passport_series", "series") or "")
        norm.setdefault("client_passport_number", first("client_passport_number", "passport_number", "number") or "")
        norm.setdefault("client_address", first("client_address", "registration_address", "address") or "")
        norm.setdefault("client_phone", first("client_phone", "phone", "tel") or "")
        norm.setdefault("email", first("email") or "")
        norm.setdefault("birth_date", first("birth_date", "dob") or "")
        norm.setdefault("birth_place", first("birth_place", "pob") or "")
        
        # Dates
        if not norm.get("contract_date"):
            norm["contract_date"] = datetime.now().strftime("%d.%m.%Y")
        
        return norm
