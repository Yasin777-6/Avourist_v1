"""
Fill DOC/DOCX templates using LibreOffice conversion + python-docx.
"""
import logging
import subprocess
import tempfile
import os
import shutil
from pathlib import Path
from typing import Dict
from datetime import datetime
from docx import Document

logger = logging.getLogger(__name__)


class DOCTextReplacer:
    """
    Replace text in DOC files by converting to DOCX first using LibreOffice.
    """
    
    def __init__(self):
        self.libreoffice_path = self._find_libreoffice()
    
    def _find_libreoffice(self) -> str:
        """Find LibreOffice executable path."""
        possible_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "soffice",  # If in PATH
        ]
        
        for path in possible_paths:
            if os.path.exists(path) or path == "soffice":
                return path
        
        logger.warning("LibreOffice not found in standard locations")
        return "soffice"
    
    def fill_doc(self, template_path: str, data: Dict, output_path: str) -> bytes:
        """
        Fill DOC template with data.
        Strategy:
        1. Try opening as DOCX (some .doc files are actually DOCX)
        2. If fails, convert .doc to .docx using LibreOffice
        3. Fill the DOCX with python-docx
        4. Save filled DOCX
        """
        template_path_obj = Path(template_path)
        output_path_obj = Path(output_path)
        
        # Determine if we need conversion
        docx_to_fill = None
        needs_conversion = False
        
        try:
            # Try opening as DOCX first
            doc = Document(template_path)
            logger.info("Template is already DOCX-compatible, filling directly...")
            docx_to_fill = template_path
            
        except Exception as e:
            logger.info(f"Template is old .doc format: {e}")
            logger.info("Converting .doc to .docx using LibreOffice...")
            needs_conversion = True
            
            # Convert using LibreOffice
            try:
                docx_to_fill = self._convert_doc_to_docx(template_path)
                logger.info(f"Converted to: {docx_to_fill}")
            except Exception as conv_error:
                logger.error(f"LibreOffice conversion failed: {conv_error}")
                raise
        
        # Now fill the DOCX file
        try:
            doc = Document(docx_to_fill)
            replacements = self._build_replacements(data)
            
            logger.info(f"Filling document with {len(replacements)} field mappings...")
            replaced_count = 0
            
            # Debug: Log paragraphs containing "Паспорт" or "Серия" to see exact text
            for i, paragraph in enumerate(doc.paragraphs):
                if "Паспорт" in paragraph.text or "Серия" in paragraph.text:
                    logger.info(f"DEBUG Passport paragraph {i}: '{paragraph.text}'")
                if "______ (______ тысяч) рублей" in paragraph.text:
                    logger.info(f"DEBUG Success fee paragraph {i}: '{paragraph.text}'")
            
            # Replace in paragraphs - handle text that spans multiple runs
            for paragraph in doc.paragraphs:
                for placeholder, replacement in replacements.items():
                    if placeholder in paragraph.text:
                        # Get full text and find placeholder
                        full_text = paragraph.text
                        if placeholder in full_text:
                            # Replace the entire paragraph text
                            new_text = full_text.replace(placeholder, str(replacement))
                            
                            # Clear all runs and add new text
                            for run in paragraph.runs:
                                run.text = ''
                            if paragraph.runs:
                                paragraph.runs[0].text = new_text
                            else:
                                paragraph.add_run(new_text)
                            
                            replaced_count += 1
                            logger.debug(f"Replaced in paragraph: '{placeholder[:30]}...'")
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, replacement in replacements.items():
                                if placeholder in paragraph.text:
                                    # Get full text and find placeholder
                                    full_text = paragraph.text
                                    if placeholder in full_text:
                                        # Replace the entire paragraph text
                                        new_text = full_text.replace(placeholder, str(replacement))
                                        
                                        # Clear all runs and add new text
                                        for run in paragraph.runs:
                                            run.text = ''
                                        if paragraph.runs:
                                            paragraph.runs[0].text = new_text
                                        else:
                                            paragraph.add_run(new_text)
                                        
                                        replaced_count += 1
                                        logger.debug(f"Replaced in table: '{placeholder[:30]}...'")
            
            logger.info(f"✓ Replaced {replaced_count} field occurrences")
            
            # Save filled document as DOCX
            output_docx = str(output_path_obj).replace('.doc', '.docx')
            doc.save(output_docx)
            
            # Read and return bytes
            with open(output_docx, 'rb') as f:
                content = f.read()
            
            logger.info(f"✓ Filled document saved: {output_docx} ({len(content)} bytes)")
            self._log_contract_data(data)
            
            # Clean up temporary converted file if needed
            if needs_conversion and docx_to_fill and os.path.exists(docx_to_fill):
                try:
                    os.remove(docx_to_fill)
                except:
                    pass
            
            return content
            
        except Exception as e:
            logger.error(f"Failed to fill document: {e}")
            raise
    
    def _convert_doc_to_docx(self, doc_path: str) -> str:
        """Convert .doc to .docx using LibreOffice."""
        doc_path_obj = Path(doc_path)
        output_dir = doc_path_obj.parent
        
        # Run LibreOffice conversion
        # Pre-check availability when using PATH reference
        if self.libreoffice_path == 'soffice' and shutil.which('soffice') is None:
            raise FileNotFoundError(
                "LibreOffice (soffice) is required to convert .doc to .docx but was not found. "
                "Install LibreOffice on the server or provide a .docx version of the template."
            )

        cmd = [
            self.libreoffice_path,
            '--headless',
            '--convert-to', 'docx',
            '--outdir', str(output_dir),
            str(doc_path)
        ]
        
        logger.info(f"Running: {' '.join(cmd)}")
        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=30
            )
        except FileNotFoundError:
            # Executable not found at provided path
            raise FileNotFoundError(
                "LibreOffice (soffice) executable not found. Install LibreOffice or switch to a .docx template."
            )
        
        if result.returncode != 0:
            raise Exception(f"LibreOffice conversion failed: {result.stderr}")
        
        # Find the converted file
        docx_path = doc_path_obj.with_suffix('.docx')
        if not docx_path.exists():
            raise Exception(f"Converted file not found: {docx_path}")
        
        return str(docx_path)
    
    def _build_replacements(self, data: Dict) -> Dict[str, str]:
        """Build dictionary of text to replace based on user's sample data."""
        replacements = {}
        
        # Contract number: "1-Б/24"
        contract_number = data.get('contract_number', '')
        if contract_number:
            replacements["1-Б/24"] = contract_number
        
        # Date: "28 апреля 2024 г" -> actual date
        contract_date = data.get('contract_date', datetime.now().strftime('%d.%m.%Y'))
        replacements["28 апреля 2024 г"] = contract_date
        
        # Client name: "Тытюк Александр Михайлович"
        client_name = data.get('client_full_name', '')
        if client_name:
            replacements["Тытюк Александр Михайлович"] = client_name
        
        # Director name (Executor side)
        director_name = data.get('director_name', 'Шельмина Евгения Васильевича')
        replacements["Шельмина Евгения Васильевича"] = director_name
        
        # Pricing section - exact patterns from your template
        total_amount = data.get('total_amount')
        if total_amount:
            amount_text = self._format_amount(int(total_amount))
            # Pattern: "составляет____ (_______ тысяч) рублей"
            replacements["составляет____ (_______ тысяч) рублей"] = f"составляет{amount_text}"
            replacements["____ (_______ тысяч) рублей"] = amount_text
            replacements["________ (________ тысяч) рублей"] = amount_text
            replacements["30 000 (тридцать тысяч) рублей"] = amount_text
            replacements["25 000 (двадцать пять тысяч) рублей"] = amount_text
        
        # Success fee - MUST come before prepayment to avoid conflicts
        # Pattern: "______ (______ тысяч) рублей исключительно при положительном решении"
        success_fee = data.get('success_fee')
        if success_fee:
            fee_text = self._format_amount(int(success_fee))
            # Use more specific pattern with context
            replacements["______ (______ тысяч) рублей исключительно при положительном решении"] = f"{fee_text} исключительно при положительном решении"
            replacements["10 000 (десять тысяч) рублей исключительно при положительном решении"] = f"{fee_text} исключительно при положительном решении"
        
        # Prepayment - multiple patterns with underscores
        prepayment = data.get('prepayment')
        if prepayment:
            prep_text = self._format_amount(int(prepayment))
            # Remove the ____ prefix by replacing the whole pattern
            replacements["____53 000 (пятьдесят три тысячи) рублей"] = prep_text
            replacements["__15 000 (пятнадцать тысяч) рублей"] = prep_text
            replacements["_______ (______ тысяч) рублей не позднее одного дня"] = f"{prep_text} не позднее одного дня"
            replacements["20 000 (двадцать тысяч) рублей"] = prep_text
            replacements["15 000 (пятнадцать тысяч) рублей"] = prep_text
        
        # Docs prep fee - pattern: "составляет5 000 (пять тысяч) рублей"
        docs_fee = data.get('docs_prep_fee')
        if docs_fee:
            docs_text = self._format_amount(int(docs_fee))
            replacements["составляет5 000 (пять тысяч) рублей"] = f"составляет{docs_text}"
            replacements["5 000 (пять тысяч) рублей"] = docs_text
            replacements[" ______(___________ тысяч) рублей"] = f" {docs_text}"
            replacements["______(___________ тысяч) рублей"] = docs_text
            replacements["5 000 рублей"] = docs_text
        
        # Client details section
        birth_date = data.get('birth_date', '')
        if birth_date:
            replacements["Дата/ месяц/ год рождения"] = birth_date
        
        birth_place = data.get('birth_place', '')
        if birth_place:
            replacements["Место рождения"] = birth_place
        
        passport_series = data.get('client_passport_series', '')
        passport_number = data.get('client_passport_number', '')
        if passport_series and passport_number:
            # Exact pattern from template: "Паспорт Серия_____  Номер___________" (note 2 spaces)
            passport_filled = f"Серия {passport_series} Номер {passport_number}"
            replacements["Паспорт Серия_____  Номер___________"] = f"Паспорт {passport_filled}"
            replacements["Серия_____  Номер___________"] = passport_filled
            # Also try with single space
            replacements["Паспорт Серия_____ Номер___________"] = f"Паспорт {passport_filled}"
            replacements["Серия_____ Номер___________"] = passport_filled
        
        address = data.get('client_address', '')
        if address:
            # Pattern with underscores
            replacements["Зарегистрирован: _____________________"] = f"Зарегистрирован: {address}"
        
        phone = data.get('client_phone', '')
        if phone:
            # Pattern with underscores
            replacements["Тел. _________________________"] = f"Тел. {phone}"
        
        email = data.get('email', '')
        if email:
            # Multiple email patterns - with and without space after colon
            replacements["Е-mail______________________________"] = f"Е-mail: {email}"
            replacements["Е-mail:______________________________"] = f"Е-mail: {email}"
            replacements["E-mail______________________________"] = f"E-mail: {email}"
        
        # Case description
        case_desc = data.get('case_description', '')
        if case_desc:
            default_desc = "провести правовой анализ документов (материалов дела), подготовитьходатайство на получение материалов дела, ходатайство о привлечении к делу защитника, ходатайство о переводе дела по месту жительства, подготовка письменного объяснения лица по делу об административном правонарушении по ч.1 ст.12.8 КоАП РФ"
            replacements[default_desc] = case_desc
        
        return replacements
    
    def _format_amount(self, amount: int) -> str:
        """Format amount as '15 000 (пятнадцать тысяч) рублей'."""
        try:
            from num2words import num2words
            words = num2words(int(amount), lang="ru")
            formatted = f"{int(amount):,}".replace(",", " ")
            return f"{formatted} ({words}) рублей"
        except Exception as e:
            logger.warning(f"Could not format amount with words: {e}")
            formatted = f"{int(amount):,}".replace(",", " ")
            return f"{formatted} рублей"
    
    def _log_contract_data(self, data: Dict):
        """Log contract data for verification."""
        logger.info("=" * 50)
        logger.info("CONTRACT DATA FILLED:")
        logger.info(f"  Contract #: {data.get('contract_number', 'N/A')}")
        logger.info(f"  Date: {data.get('contract_date', 'N/A')}")
        logger.info(f"  Client: {data.get('client_full_name', 'N/A')}")
        logger.info(f"  Birth: {data.get('birth_date', 'N/A')} at {data.get('birth_place', 'N/A')}")
        logger.info(f"  Passport: {data.get('client_passport_series', 'N/A')} {data.get('client_passport_number', 'N/A')}")
        logger.info(f"  Address: {data.get('client_address', 'N/A')}")
        logger.info(f"  Phone: {data.get('client_phone', 'N/A')}")
        logger.info(f"  Email: {data.get('email', 'N/A')}")
        logger.info(f"  Total: {data.get('total_amount', 'N/A')} руб.")
        logger.info(f"  Prepayment: {data.get('prepayment', 'N/A')} руб.")
        logger.info(f"  Success Fee: {data.get('success_fee', 'N/A')} руб.")
        logger.info(f"  Docs Fee: {data.get('docs_prep_fee', 'N/A')} руб.")
        logger.info("=" * 50)
