"""
Create a DOCX template with placeholder text for contract generation.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create new document
doc = Document()

# Add title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Договор № ___ об оказании юридических услуг")
run.bold = True
run.font.size = Pt(14)

# Add date and location
date_para = doc.add_paragraph()
date_para.add_run("г. Москва                                                    ")
date_run = date_para.add_run("28 апреля 2024 г")
date_run.font.color.rgb = RGBColor(255, 0, 0)  # Red

# Add parties
parties = doc.add_paragraph()
name_run = parties.add_run("Тытюк Александр Михайлович")
name_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
parties.add_run(" (именуемый в дальнейшем «Заказчик») и Автономная некоммерческая организация по предоставлению услуг в области права «Первое Арбитражное Учреждение» в лице директора Шельмина Евгения Васильевича, (именуемый в дальнейшем «Исполнитель»), с другой стороны, именуемые в дальнейшем «Стороны», заключили настоящий Договор о нижеследующем:")

# Section 1
doc.add_heading("1. Предмет договора", level=2)
section1 = doc.add_paragraph()
section1.add_run("1.1 Заказчик поручает, а Исполнитель принимает на себя обязательства оказать юридические услуги: провести правовой анализ документов (материалов дела), ")
case_desc_run = section1.add_run("подготовка ответа на требование каршеринга, подготовка претензии, подготовка Отзыва на исковое заявления, ответ на претензию, заявление в соответствующие органы")
case_desc_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
section1.add_run(". Исполнитель представляет интересы Заказчика в суде и в иных органах по нотариальной доверенности без личного присутствия.")

# Section 2
doc.add_heading("2. Обязанности и права сторон", level=2)
doc.add_paragraph("2.1. Заказчик обязан:\nа) произвести оплату по договору согласно п. 5.1 до момента начала оказания Исполнителем Услуг.\nб) обеспечить своевременное предоставление Исполнителю всей информации и документации, необходимой для оказания Услуг.")

# Section 5 - Pricing
doc.add_heading("5. Стоимость услуг и порядок расчетов", level=2)

pricing_para = doc.add_paragraph()
pricing_para.add_run("5.1. Стоимость услуг по договору в соответствии с п.1.1. настоящего Договора составляет ")
total_run = pricing_para.add_run("_________(______________)  рублей")
total_run.font.color.rgb = RGBColor(255, 0, 0)
pricing_para.add_run(". Оплата производится в следующем порядке: ")
prep_run = pricing_para.add_run("____________ (_________ тысяч) рублей")
prep_run.font.color.rgb = RGBColor(255, 0, 0)
pricing_para.add_run(" не позднее одного дня с даты подписания настоящего Договора Сторонами. ")
success_run = pricing_para.add_run("__________________ (___________ тысяч) рублей")
success_run.font.color.rgb = RGBColor(255, 0, 0)
pricing_para.add_run(" исключительно при положительном решении в соответствии с п.8.8 настоящего Договора.")

doc.add_paragraph()
docs_para = doc.add_paragraph()
docs_para.add_run("5.1.1 Подготовка документов в соответствии с п.1.1. настоящего Договора составляет ")
docs_run = docs_para.add_run("______ (___________ тысяч) рублей")
docs_run.font.color.rgb = RGBColor(255, 0, 0)
docs_para.add_run(".")

# Add more sections...
doc.add_heading("9. Адреса, реквизиты и подписи Сторон:", level=2)

# Create table for signatures
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Left column - Executor
left_cell = table.rows[0].cells[0]
left_cell.text = "Исполнитель\n\nАНО «Первое Арбитражное Учреждение»\n\n_______________________\nДиректор Шельмин Е.В."

# Right column - Client
right_cell = table.rows[0].cells[1]
client_info = right_cell.add_paragraph()
client_name_run = client_info.add_run("Тытюк Александр Михайлович")
client_name_run.font.color.rgb = RGBColor(255, 0, 0)

client_info.add_run("\n\nОБЯЗАТЕЛЬНО !!!!!!!!!!!!!!!!!!!\n")

birth_date_run = client_info.add_run("Дата/ месяц/ год рождения")
birth_date_run.font.color.rgb = RGBColor(255, 0, 0)

client_info.add_run("\n")

birth_place_run = client_info.add_run("Место рождения")
birth_place_run.font.color.rgb = RGBColor(255, 0, 0)

client_info.add_run("\n")

passport_run = client_info.add_run("Паспорт Серия_____ Номер___________")
passport_run.font.color.rgb = RGBColor(255, 0, 0)

client_info.add_run("\n\n")

address_run = client_info.add_run("Зарегистрирован: _____________________")
address_run.font.color.rgb = RGBColor(255, 0, 0)

client_info.add_run("\n\n")

phone_run = client_info.add_run("Тел. _________________________")
phone_run.font.color.rgb = RGBColor(255, 0, 0)

client_info.add_run("\n\n")

email_run = client_info.add_run("Е-mail______________________________")
email_run.font.color.rgb = RGBColor(255, 0, 0)

client_info.add_run("\n\n")

signature_run = client_info.add_run("_______________________________________\n(Фамилия И.О прописью / подпись)")
signature_run.font.color.rgb = RGBColor(255, 0, 0)

# Save template
output_path = "contracts/contract_template.docx"
doc.save(output_path)
print(f"✓ Created DOCX template: {output_path}")
